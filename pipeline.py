# pipeline.py

from dataclasses import dataclass
from pathlib import Path
import shutil
import uuid
import zipfile

from docx import Document
from PyPDF2 import PdfReader
import fitz

from config import UPLOADS_DIR
from ocr.quality import analyse_image
from ocr.preprocessor import preprocess_image
from ocr.engines import get_ocr_engine

from ai.extractor import get_ai_extractor

from update_employee import merge_extraction_results, map_extracted_to_employee_fields, save_employee_draft

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".heic"}
PDF_EXTENSIONS   = {".pdf"}
DOCX_EXTENSIONS  = {".docx"}

MIN_USEFUL_OCR_CHARS = 40


# ── Upload session ────────────────────────────────────────────────────────────

@dataclass
class UploadSession:
    """
    Holds all directory paths for a single upload run.
    Every file produced during that run lives inside session_dir so uploads
    stay isolated and traceable (upload_0001/, upload_0002/, …).
    """
    name: str               # e.g. "upload_0003"
    session_dir: Path
    originals_dir: Path
    preprocessed_dir: Path
    extracted_images_dir: Path
    debug_dir: Path
    ocr_texts_dir: Path


def _create_upload_session() -> UploadSession:
    """
    Scan uploads/ for existing upload_NNNN folders, then create the next one.
    All four subdirectories are created up-front so callers never need to mkdir.
    """
    existing_nums = []
    for d in UPLOADS_DIR.iterdir():
        if d.is_dir() and d.name.startswith("upload_"):
            try:
                existing_nums.append(int(d.name.split("_")[1]))
            except (IndexError, ValueError):
                pass

    next_num = max(existing_nums, default=0) + 1
    name = f"upload_{next_num:04d}"
    session_dir = UPLOADS_DIR / name

    session = UploadSession(
        name=name,
        session_dir=session_dir,
        originals_dir=session_dir / "originals",
        preprocessed_dir=session_dir / "preprocessed",
        extracted_images_dir=session_dir / "extracted_images",
        debug_dir=session_dir / "debug",
        ocr_texts_dir=session_dir / "ocr_texts",
    )

    for d in [session.originals_dir, session.preprocessed_dir,
              session.extracted_images_dir, session.debug_dir,
              session.ocr_texts_dir]:
        d.mkdir(parents=True, exist_ok=True)

    return session


# ── File saving ───────────────────────────────────────────────────────────────

def save_uploaded_file(upload_file, session: UploadSession) -> str:
    suffix = Path(upload_file.filename).suffix.lower()
    safe_name = f"{uuid.uuid4().hex}{suffix}"
    target = session.originals_dir / safe_name
    with target.open("wb") as buffer:
        shutil.copyfileobj(upload_file.file, buffer)
    return str(target)


# ── Text extraction helpers ───────────────────────────────────────────────────

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)
    return "\n".join(paragraphs + table_texts).strip()


def extract_text_from_pdf_if_possible(pdf_path: str) -> str:
    try:
        reader = PdfReader(pdf_path)
        texts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(t for t in texts if t.strip()).strip()
    except Exception:
        return ""


# ── Image extraction helpers ──────────────────────────────────────────────────

def extract_images_from_docx(docx_path: str, session: UploadSession) -> list[str]:
    output_paths = []
    extract_dir = session.extracted_images_dir / Path(docx_path).stem
    extract_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                filename = Path(name).name
                out_path = extract_dir / filename
                with z.open(name) as src, open(out_path, "wb") as dst:
                    dst.write(src.read())
                output_paths.append(str(out_path))

    return output_paths


def render_pdf_pages_to_images(pdf_path: str, session: UploadSession) -> list[str]:
    """Render each PDF page to PNG. Used only when the PDF has no embedded images."""
    doc = fitz.open(pdf_path)
    output_dir = session.extracted_images_dir / f"{Path(pdf_path).stem}_pages"
    output_dir.mkdir(parents=True, exist_ok=True)

    image_paths = []
    for page_index in range(len(doc)):
        page = doc[page_index]
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
        out_path = output_dir / f"page_{page_index + 1}.png"
        pix.save(str(out_path))
        image_paths.append(str(out_path))

    doc.close()
    return image_paths


def extract_embedded_images_from_pdf(pdf_path: str, session: UploadSession) -> list[str]:
    """Extract raw embedded images from a scanned PDF — preferred over page rendering."""
    doc = fitz.open(pdf_path)
    output_dir = session.extracted_images_dir / f"{Path(pdf_path).stem}_embedded"
    output_dir.mkdir(parents=True, exist_ok=True)

    image_paths = []
    for page_index in range(len(doc)):
        page = doc[page_index]
        for image_index, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]
            base_image = doc.extract_image(xref)
            ext = base_image["ext"]
            out_path = output_dir / f"page_{page_index + 1}_img_{image_index}.{ext}"
            with open(out_path, "wb") as f:
                f.write(base_image["image"])
            image_paths.append(str(out_path))

    doc.close()
    return image_paths


# ── OCR scoring ───────────────────────────────────────────────────────────────

def score_ocr_text(text: str) -> int:
    """
    Score OCR text quality. Higher = more useful.
    Rewards real letter/digit content and known German form keywords.
    Penalises outputs that are mostly garbage (low alpha ratio).
    """
    if not text:
        return 0

    total = len(text)
    alpha = sum(ch.isalpha() for ch in text)
    digits = sum(ch.isdigit() for ch in text)

    # If fewer than 20% of characters are letters the output is likely garbled
    alpha_ratio = alpha / total if total else 0
    if alpha_ratio < 0.20:
        return max(0, total // 4)

    score = total + alpha + digits

    keywords = [
        "name", "vorname", "familienname", "geburtsdatum", "geburtsort",
        "nationality", "staatsangehörigkeit", "telefon", "iban", "bic",
        "ausweis", "reisepass", "aufenthalt", "steuer", "versicherung",
    ]
    lowered = text.lower()
    for kw in keywords:
        if kw in lowered:
            score += 25

    return score


# ── Document helpers ─────────────────────────────────────────────────────────

def _is_mrz_only(text: str) -> bool:
    """
    Return True when the page is a MRZ-only page (back of ID/residence permit).
    Uses line-level pattern matching: any line with 5+ consecutive < or > chars
    is a strong MRZ signal, regardless of how much noise text surrounds it.
    """
    if not text:
        return False
    import re
    for line in text.splitlines():
        if re.search(r"[<>]{5,}", line):
            return True
    return False


def _detect_doc_type(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ("reisepass", "passport", "p<ukr", "p<deu", "p<afg", "p<pol")):
        return "passport"
    if any(k in t for k in ("aufenthaltstitel", "aufenthaltserlaubnis", "residence permit",
                             "ausweis für vertriebene", "ausweis fur vertriebene")):
        return "residence_permit"
    if any(k in t for k in ("melderegister", "meldezettel", "bestätigung der meldung",
                             "bestatigung der meldung", "zentrales melderegister")):
        return "meldezettel"
    if any(k in t for k in ("iban", "bic", "kontonummer", "kreditkarte", "sparkasse",
                             "bank austria", "erste bank", "raiffeisen")):
        return "bank_document"
    if any(k in t for k in ("sozialversicherung", "sv-nummer", "e-card")):
        return "sv_card"
    if any(k in t for k in ("krankenkasse", "krankenversicherung", "versicherungskarte")):
        return "health_card"
    if any(k in t for k in ("steuer-id", "steuernummer", "finanzamt")):
        return "tax_document"
    return "document"


def _save_ocr_text(session: UploadSession, stem: str, original_name: str, text: str) -> None:
    out = session.ocr_texts_dir / f"{stem}.txt"
    out.write_text(f"# Source: {original_name}\n\n{text}", encoding="utf-8")


# ── Core image processing ─────────────────────────────────────────────────────

def process_image_file(image_path: str, session: UploadSession) -> dict:
    quality = analyse_image(image_path)

    if quality.width < 300 or quality.height < 150:
        raise ValueError(
            f"Image too small for useful OCR: {image_path} ({quality.width}x{quality.height})"
        )

    prep = preprocess_image(
        image_path,
        preprocessed_dir=session.preprocessed_dir,
        debug_dir=session.debug_dir,
    )
    final_image = prep["final_path"]

    if not Path(final_image).exists():
        raise FileNotFoundError(f"Preprocessed image was not created: {final_image}")

    ocr = get_ocr_engine("tesseract")
    ocr_trials = []

    for candidate_name, candidate_path in prep.get("ocr_candidates", {}).items():
        if not candidate_path or not Path(candidate_path).exists():
            continue
        try:
            candidate_text = ocr.extract_text(candidate_path)
            ocr_trials.append({
                "candidate_name": candidate_name,
                "candidate_path": candidate_path,
                "raw_text": candidate_text,
                "score": score_ocr_text(candidate_text),
            })
        except Exception:
            continue

    if not ocr_trials:
        raise ValueError(
            f"OCR failed for all preprocessing candidates for image: {image_path}."
        )

    best_trial = max(ocr_trials, key=lambda x: x["score"])
    raw_text = best_trial["raw_text"]
    best_candidate_name = best_trial["candidate_name"]
    best_candidate_path = best_trial["candidate_path"]

    stem = Path(image_path).stem
    _save_ocr_text(session, stem, Path(image_path).name, raw_text)

    # Skip AI extraction when OCR output is too sparse — blank backs of documents,
    # very dark images, or wholly unreadable scans.
    if len(raw_text.strip()) < MIN_USEFUL_OCR_CHARS:
        print(f"[INFO] Skipping AI extraction for {Path(image_path).name}: "
              f"OCR text too short ({len(raw_text.strip())} chars)")
        return {
            "source": image_path,
            "raw_text": raw_text,
            "extracted_data": {},
            "best_candidate_name": best_candidate_name,
            "best_candidate_path": best_candidate_path,
            "ocr_trials": ocr_trials,
            "debug_steps": prep["steps"],
            "pipeline_path": "skipped_insufficient_ocr_text",
        }

    # Skip when the page is pure MRZ (back of ID/permit) — all <> filler, no real data.
    if _is_mrz_only(raw_text):
        print(f"[INFO] Skipping AI extraction for {Path(image_path).name}: MRZ-only page")
        return {
            "source": image_path,
            "raw_text": raw_text,
            "extracted_data": {},
            "best_candidate_name": best_candidate_name,
            "best_candidate_path": best_candidate_path,
            "ocr_trials": ocr_trials,
            "debug_steps": prep["steps"],
            "pipeline_path": "skipped_mrz_only",
        }

    doc_type = _detect_doc_type(raw_text)
    extractor = get_ai_extractor("local")
    extracted = extractor.extract_from_text(
        raw_text,
        source_description=f"Document type: {doc_type}",
    )

    return {
        "source": image_path,
        "quality": {
            "width": quality.width,
            "height": quality.height,
            "brightness_score": quality.brightness_score,
            "blur_score": quality.blur_score,
            "is_dark": quality.is_dark,
            "is_blurry": quality.is_blurry,
            "is_small": quality.is_small,
            "needs_preprocessing": quality.needs_preprocessing,
            "route_to_vision": quality.route_to_vision,
        },
        "best_candidate_name": best_candidate_name,
        "best_candidate_path": best_candidate_path,
        "ocr_trials": ocr_trials,
        "raw_text": raw_text,
        "extracted_data": extracted,
        "debug_steps": prep["steps"],
        "pipeline_path": "image_preprocess_ocr_local_extract",
    }


# ── Document-type dispatchers ─────────────────────────────────────────────────

def process_docx_file(docx_path: str, session: UploadSession) -> dict:
    direct_text = extract_text_from_docx(docx_path)
    image_paths = extract_images_from_docx(docx_path, session)
    results = []

    if direct_text.strip():
        extractor = get_ai_extractor("local")
        extracted = extractor.extract_from_text(
            direct_text,
            source_description=f"DOCX text: {Path(docx_path).name}",
        )
        results.append({
            "source": f"{docx_path}::text",
            "raw_text": direct_text,
            "extracted_data": extracted,
            "pipeline_path": "docx_text_local_extract",
        })

    for img_path in image_paths:
        results.append(process_image_file(img_path, session))

    merged = merge_extraction_results(results)
    return {
        "source": docx_path,
        "raw_text": direct_text,
        "embedded_images": image_paths,
        "sub_results": results,
        "extracted_data": merged,
        "verification": {"overall_confidence": "medium", "ready_for_review": True, "issues": []},
        "pipeline_path": "docx_mixed_content_pipeline",
    }


def process_pdf_file(pdf_path: str, session: UploadSession) -> dict:
    direct_text = extract_text_from_pdf_if_possible(pdf_path)
    results = []

    if direct_text.strip():
        extractor = get_ai_extractor("local")
        extracted = extractor.extract_from_text(
            direct_text,
            source_description=f"PDF text layer: {Path(pdf_path).name}",
        )
        results.append({
            "source": f"{pdf_path}::text",
            "raw_text": direct_text,
            "extracted_data": extracted,
            "pipeline_path": "pdf_text_local_extract",
        })

    # Prefer embedded images (original scan quality) over page rendering.
    # Only render pages when the PDF has no embedded images (digital/vector PDF).
    embedded_images: list[str] = []
    page_images: list[str] = []

    try:
        embedded_images = extract_embedded_images_from_pdf(pdf_path, session)
    except Exception:
        embedded_images = []

    if not embedded_images:
        try:
            page_images = render_pdf_pages_to_images(pdf_path, session)
        except Exception:
            page_images = []

    for emb_img in embedded_images:
        try:
            results.append(process_image_file(emb_img, session))
        except Exception as e:
            print(f"[WARN] Skipping embedded PDF image {emb_img}: {e}")

    for page_img in page_images:
        try:
            results.append(process_image_file(page_img, session))
        except Exception as e:
            print(f"[WARN] Skipping rendered PDF page image {page_img}: {e}")

    if not results:
        raise ValueError(
            "PDF could not be processed. No text layer found and image OCR path failed."
        )

    merged = merge_extraction_results(results)
    return {
        "source": pdf_path,
        "raw_text": direct_text,
        "page_images": page_images,
        "embedded_images": embedded_images,
        "sub_results": results,
        "extracted_data": merged,
        "verification": {"overall_confidence": "medium", "ready_for_review": True, "issues": []},
        "pipeline_path": "pdf_mixed_content_pipeline",
    }


def process_single_file(file_path: str, session: UploadSession) -> dict:
    suffix = Path(file_path).suffix.lower()

    if suffix in IMAGE_EXTENSIONS:
        return process_image_file(file_path, session)
    if suffix in DOCX_EXTENSIONS:
        return process_docx_file(file_path, session)
    if suffix in PDF_EXTENSIONS:
        return process_pdf_file(file_path, session)

    raise ValueError(f"Unsupported file type: {suffix}")


# ── Entry point ───────────────────────────────────────────────────────────────

def process_uploaded_files(
    db,
    files=None,
    text_input: str | None = None,
    employee_id: int | None = None,
):
    files = files or []
    session = _create_upload_session()
    print(f"[INFO] Upload session: {session.name} ({session.session_dir})")

    all_results = []
    ocr_report_sections = []

    for upload_file in files:
        if not upload_file or not upload_file.filename:
            continue
        suffix = Path(upload_file.filename).suffix.lower()
        if not suffix:
            continue

        saved_path = save_uploaded_file(upload_file, session)
        result = process_single_file(saved_path, session)
        all_results.append(result)

        # Collect OCR text for the session report
        raw = result.get("raw_text") or ""
        sub_texts = [
            s.get("raw_text", "") for s in result.get("sub_results", [])
            if s.get("raw_text") and s.get("raw_text") != raw
        ]
        all_text = "\n---\n".join([raw] + sub_texts) if raw else "(no text extracted)"
        doc_type = _detect_doc_type(raw)
        ocr_report_sections.append(
            f"=== {upload_file.filename} [{doc_type}] ===\n{all_text}"
        )

    if text_input and text_input.strip():
        extractor = get_ai_extractor("local")
        extracted = extractor.extract_from_text(text_input, source_description="Pasted text")
        all_results.append({
            "source": "text_input",
            "raw_text": text_input,
            "extracted_data": extracted,
            "pipeline_path": "text_local_extract",
        })

    if not all_results:
        raise ValueError("No input provided")

    # Write OCR report — one file per session showing what text was extracted from each document
    if ocr_report_sections:
        report_path = session.session_dir / "ocr_report.txt"
        report_path.write_text("\n\n".join(ocr_report_sections), encoding="utf-8")
        print(f"[INFO] OCR report saved: {report_path}")

    merged = merge_extraction_results(all_results)
    mapped = map_extracted_to_employee_fields(merged)
    employee = save_employee_draft(db, mapped, employee_id=employee_id)

    return {
        "employee_id": employee.id,
        "employee": employee,
        "session_name": session.name,
        "session_dir": str(session.session_dir),
        "results": all_results,
        "merged_data": merged,
        "mapped_data": mapped,
    }
